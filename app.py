from flask import Flask, render_template, url_for, request, flash, redirect, send_file, session
import os
import re
from werkzeug.utils import secure_filename
from docx import Document
from reportlab.pdfgen import canvas
from docx2pdf import convert
from fpdf import FPDF
from os.path import join, exists
from PyPDF2 import PdfReader
from docx import Document
from flask import request

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'docx'}

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = 'supersecretkey'

def allowed_file(filename):

    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_file_extension(filename):

    _, extension = os.path.splitext(filename)
    return extension.lower()

def remove_existing_file():

    current_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'current_file')
    for extension in ALLOWED_EXTENSIONS:
        if os.path.exists(current_file_path + '.' + extension):
            os.remove(current_file_path + '.' + extension)

def extract_text_from_docx(docx_path):

    doc = Document(docx_path)
    return ' '.join([paragraph.text for paragraph in doc.paragraphs])

def extract_text_from_file(file_path):

    _, extension = os.path.splitext(file_path)
    
    if extension.lower() == '.docx':
        return extract_text_from_docx(file_path)
    elif extension.lower() == '.txt':
        return extract_text_from_txt(file_path)
    else:
        # Handle other file types if needed
        return None

def extract_text_from_txt(txt_path):

    with open(txt_path, 'r', encoding='utf-8') as txt_file:
        return txt_file.read()

def generate_filled_pdf(new_text, file_path, extension):

    print("Generating PDF...")
    print("File Path:", file_path)
    print("Extension:", extension)
    pdf_filename = os.path.join(app.config['UPLOAD_FOLDER'], 'generated_file.pdf')

    if extension == '.docx':
        # Convert DOCX to PDF using docx2pdf
        convert(file_path, pdf_filename)
    elif extension == '.txt':
        # Create a PDF document using fpdf
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, new_text)
        pdf.output(pdf_filename)

    return pdf_filename

def get_passages(file_content):

    passages = []
    current_passage = ""
    for char in file_content:
        if char == '_':
            current_passage += char
        else:
            if current_passage:
                passages.append(current_passage)
                current_passage = ""
    if current_passage:
        passages.append(current_passage)
    return passages

def count_underscores(text):

    passages = get_passages(text)
    return len(passages)

@app.route('/', methods=['GET', 'POST'])
def index():

    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)

        # Remove existing file before saving the new one
        remove_existing_file()

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'current_file' + os.path.splitext(filename)[1])
            file.save(file_path)
            flash('File uploaded successfully')

            # Extract passages from the uploaded file
            file_content = extract_text_from_file(file_path)
            passages = get_passages(file_content)

            # Store passages in the session
            session['passages'] = passages

            # Redirect to the 'index' route
            return redirect(url_for('index'))

        if 'pdf_to_docx' in request.form: # Adăugăm condiția pentru conversia PDF to Word
            if 'current_file.pdf' in os.listdir(app.config['UPLOAD_FOLDER']):
                pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], 'current_file.pdf')
                docx_filename = convert_pdf_to_docx(pdf_path)
                return send_file(docx_filename, as_attachment=True, download_name='converted_file.docx')
            else:
                flash('Nu există un fișier PDF pentru conversie')
                return redirect(url_for('index'))

    # Renderea șablonului index.html cu modificările pentru conversia PDF to Word
    return render_template('index.html', current_file=get_current_file())


def get_current_file():

    current_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'current_file')
    for extension in ALLOWED_EXTENSIONS:
        if os.path.exists(current_file_path + '.' + extension):
            return f'Current file selected: {os.path.basename(current_file_path + "." + extension)}'
    return '[none]'


def convert_pdf_to_docx(pdf_path):
    import PyPDF2
    from docx import Document

    def extract_text_from_pdf(pdf_path):
        with open(pdf_path, 'rb') as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            num_pages = len(reader.pages)
            text = ''
            for page_num in range(num_pages):
                page = reader.getPage(page_num)
                text += page.extract_text()
            return text

    pdf_text = extract_text_from_pdf(pdf_path)

    doc = Document()
    doc.add_paragraph(pdf_text)
    docx_filename = os.path.splitext(pdf_path)[0] + '.docx'  # Schimbarea extensiei
    doc.save(docx_filename)

    return docx_filename


@app.route('/process_passage', methods=['POST'])
def process_passage():

    print("Entered process function @@@@@@@@@@@@@@@@@@@@@@@@")

    # Update with the actual file extension
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'current_file')

    # Check if either .txt or .docx file exists in the upload folder
    txt_file_exists = exists(file_path + '.txt')
    docx_file_exists = exists(file_path + '.docx')

    if not (txt_file_exists or docx_file_exists):
        flash('No valid file found')
        print("return from redirect#############################@@@@@")
        return redirect(url_for('index'))

    # Choose the first available file (either .txt or .docx)
    if txt_file_exists:
        extension = '.txt'
    else:
        extension = '.docx'

    # Get the correct file path
    file_path = file_path + extension

    if extension == '.docx':
        initial_text = extract_text_from_docx(file_path)
    elif extension == '.txt':
        initial_text = extract_text_from_txt(file_path)
    else:
        flash('Unsupported file type')
        print("Returned with the extension: @" + extension + "@")
        return redirect(url_for('index'))

    # Check if the form fields are available in the request
    textfield_inputs = []
    print ("LEN OF REQUEST FIELD IS: @" + str(len(request.form)) + "@")

    for i in range(1, len(request.form)):    # inainte era  len(request.form) + 1
        key = f'textfield_{i}'
        if key in request.form:
            textfield_inputs.append(request.form[key])
        else:
            flash(f'Missing form field: {key}')
            print("return from redirect#############################@@@@@")
            return redirect(url_for('index'))

    # Create newText by replacing ____________ with textfield_inputs
    new_text = re.sub(r'_+', '{}', initial_text).format(*textfield_inputs)

    print("@@@@@@" + new_text + "@@@@@@@@")

    pdf_filename = generate_filled_pdf(new_text, file_path, extension)
    return send_file(pdf_filename, as_attachment=True, download_name='generated_file.pdf')

@app.route('/pdf_to_docx', methods=['GET', 'POST'])
def pdf_to_docx():
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)

        file = request.files['file']
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)

        # Verifică dacă fișierul este PDF
        if file and file.filename.endswith('.pdf'):
            pdf_reader = PdfReader(file)
            pdf_text = ''
            for page in pdf_reader.pages:
                pdf_text += page.extract_text()

            doc = Document()
            doc.add_paragraph(pdf_text)
            docx_filename = 'converted_file.docx'
            doc.save(docx_filename)
            
            return send_file(docx_filename, as_attachment=True, download_name='converted_file.docx')

    return render_template('index.html')

if __name__ == "__main__":
    
    app.run(debug=True)
